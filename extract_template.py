#!/usr/bin/env python3
"""
提取Word文档内容
"""

import os
from docx import Document

def extract_word_content(file_path):
    """提取Word文档内容"""
    try:
        # 打开文档
        doc = Document(file_path)
        
        # 输出文档信息
        print("=== 文档信息 ===")
        print(f"文件路径: {file_path}")
        print(f"段落数量: {len(doc.paragraphs)}")
        print(f"表格数量: {len(doc.tables)}")
        print(f"图片数量: {len(doc.inline_shapes) + len(doc.shapes)}")
        print()
        
        # 输出所有段落
        print("=== 文档内容 ===")
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                print(f"段落 {i+1}: {para.text}")
        
        # 输出所有表格
        for i, table in enumerate(doc.tables):
            print(f"\n=== 表格 {i+1} ===")
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells]
                print("\t".join(row_text))
        
        return True
        
    except Exception as e:
        print(f"提取文档内容时出错: {e}")
        return False

if __name__ == "__main__":
    # 文档路径
    doc_path = '/root/Image Restoration/blur_image_processor/图像处理课程设计模板 本科 20230417.doc'
    
    # 检查文件是否存在
    if os.path.exists(doc_path):
        extract_word_content(doc_path)
    else:
        print(f"文件不存在: {doc_path}")
